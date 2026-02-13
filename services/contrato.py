import os
import re
from docx import Document


# -----------------------------
# Utilitários de formatação
# -----------------------------
def formatar_cnpj(cnpj: str) -> str:
    """
    Recebe CNPJ com ou sem máscara e devolve 00.000.000/0000-00
    """
    dig = re.sub(r"\D", "", cnpj or "")
    if len(dig) != 14:
        return cnpj or ""
    return f"{dig[0:2]}.{dig[2:5]}.{dig[5:8]}/{dig[8:12]}-{dig[12:14]}"


def formatar_cep(cep: str) -> str:
    """
    00000000 -> 00000-000 (se tiver 8 dígitos)
    """
    dig = re.sub(r"\D", "", cep or "")
    if len(dig) != 8:
        return cep or ""
    return f"{dig[:5]}-{dig[5:]}"


def forma_representacao_por_natureza(natureza: str) -> str:
    n = (natureza or "").upper()

    # LTDA / LIMITADA
    if "LIMITADA" in n or "LTDA" in n:
        return "Contrato Social"

    # SOCIEDADE ANÔNIMA / S.A
    if "SOCIEDADE ANONIMA" in n or "SOCIEDADE ANÔNIMA" in n or " S.A" in n or " S/A" in n or n.strip().endswith("S.A") or n.strip().endswith("S/A"):
        return "Estatuto Social"

    return "Contrato ou Estatuto Social"


def limpar_none(v):
    return (v or "").strip()


def complemento_formatado(complemento: str) -> str:
    c = limpar_none(complemento)
    # você pode escolher: retornar vazio OU algo como "Complemento: X"
    # aqui: retorna o próprio texto, mas só aparece se existir.
    return c


def montar_endereco(dados: dict) -> dict:
    """
    BrasilAPI pode variar um pouco os nomes, então tratamos os mais comuns.
    Retorna campos separados e também um ENDERECO_COMPLETO opcional.
    """
    logradouro = limpar_none(dados.get("logradouro") or dados.get("street"))
    numero = limpar_none(dados.get("numero") or dados.get("number"))
    complemento = complemento_formatado(dados.get("complemento") or dados.get("complement") or "")
    cep = formatar_cep(limpar_none(dados.get("cep") or ""))
    cidade = limpar_none(dados.get("municipio") or dados.get("cidade") or dados.get("city") or "")
    uf = limpar_none(dados.get("uf") or dados.get("estado") or dados.get("state") or "")

    # Se quiser um placeholder único de endereço completo:
    # "Rua X, nº 10, Apto 2, CEP 00000-000, Cidade/UF"
    partes = []
    if logradouro:
        partes.append(logradouro)
    if numero:
        partes.append(f"nº {numero}")
    if complemento:
        partes.append(complemento)
    if cep:
        partes.append(f"CEP {cep}")
    if cidade or uf:
        partes.append(f"{cidade}/{uf}".strip("/"))

    endereco_completo = ", ".join([p for p in partes if p])

    return {
        "LOGRADOURO": logradouro,
        "NUMERO": numero,
        "COMPLEMENTO": complemento,
        "CEP": cep,
        "CIDADE": cidade,
        "UF": uf,
        "ENDERECO_COMPLETO": endereco_completo,
    }


# -----------------------------
# Substituição no DOCX
# -----------------------------
def _replace_in_paragraph(paragraph, subs: dict):
    """
    Substitui placeholders em um parágrafo.
    Funciona melhor quando o placeholder NÃO está quebrado em múltiplos runs.
    """
    if not paragraph.runs:
        return

    # Checa rápido se tem algum placeholder no texto do parágrafo
    full_text = paragraph.text
    if "<<" not in full_text:
        return

    changed = False
    for run in paragraph.runs:
        for k, v in subs.items():
            if k in run.text:
                run.text = run.text.replace(k, v)
                changed = True

    # fallback leve: se ainda existir placeholder no texto (quebrado em runs),
    # não tentamos reconstruir (porque perderia formatação). Melhor é ajustar o Word.
    return changed


def _replace_in_table(table, subs: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, subs)


def _replace_everywhere(doc: Document, subs: dict):
    # Corpo
    for p in doc.paragraphs:
        _replace_in_paragraph(p, subs)

    # Tabelas do corpo
    for t in doc.tables:
        _replace_in_table(t, subs)

    # Cabeçalhos/rodapés
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            _replace_in_paragraph(p, subs)
        for t in header.tables:
            _replace_in_table(t, subs)

        for p in footer.paragraphs:
            _replace_in_paragraph(p, subs)
        for t in footer.tables:
            _replace_in_table(t, subs)


# -----------------------------
# API pública
# -----------------------------
def gerar_numero_contrato(numero_manual: str) -> str:
    """
    Agora o número é manual no sistema.
    Mantive a função para centralizar validações se quiser depois.
    """
    n = (numero_manual or "").strip()
    if not n:
        raise ValueError("Número do contrato é obrigatório.")
    return n


def gerar_contrato(dados_fornecedor: dict, numero_contrato: str, template_path: str) -> str:
    """
    Gera contrato a partir de um modelo (docx) e salva em /contratos.

    Placeholders esperados no Word:
      <<NUMERO_CONTRATO>>
      <<RAZAO_SOCIAL>>
      <<NOME_FANTASIA>>
      <<NATUREZA_JURIDICA>>
      <<CNPJ_FORMATADO>>
      <<LOGRADOURO>>
      <<NUMERO>>
      <<COMPLEMENTO>>
      <<CEP>>
      <<CIDADE>>
      <<UF>>
      <<FORMA_REPRESENTACAO>>

    (Opcional) <<ENDERECO_COMPLETO>>
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Modelo não encontrado: {template_path}")

    doc = Document(template_path)

    # Campos principais (BrasilAPI)
    razao_social = limpar_none(dados_fornecedor.get("razao_social"))
    nome_fantasia = limpar_none(dados_fornecedor.get("nome_fantasia"))
    natureza = limpar_none(dados_fornecedor.get("natureza_juridica"))
    cnpj_bruto = limpar_none(dados_fornecedor.get("cnpj"))

    cnpj_formatado = formatar_cnpj(cnpj_bruto)
    forma_rep = forma_representacao_por_natureza(natureza)
    end = montar_endereco(dados_fornecedor)

    subs = {
        "<<NUMERO_CONTRATO>>": limpar_none(numero_contrato),
        "<<RAZAO_SOCIAL>>": razao_social,
        "<<NOME_FANTASIA>>": nome_fantasia,
        "<<NATUREZA_JURIDICA>>": natureza,
        "<<CNPJ_FORMATADO>>": cnpj_formatado,

        "<<LOGRADOURO>>": end["LOGRADOURO"],
        "<<NUMERO>>": end["NUMERO"],
        "<<COMPLEMENTO>>": end["COMPLEMENTO"],
        "<<CEP>>": end["CEP"],
        "<<CIDADE>>": end["CIDADE"],
        "<<UF>>": end["UF"],
        "<<ENDERECO_COMPLETO>>": end["ENDERECO_COMPLETO"],

        "<<FORMA_REPRESENTACAO>>": forma_rep,
    }

    _replace_everywhere(doc, subs)

    os.makedirs("contratos", exist_ok=True)
    safe_num = re.sub(r"[^A-Za-z0-9._-]+", "_", numero_contrato.strip())
    output_path = os.path.join("contratos", f"{safe_num}.docx")
    doc.save(output_path)
    return output_path
